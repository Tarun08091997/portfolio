"""
Resume Generator for Tarun Kumar - Data Science Resume
Generates a Word document resume based on project and experience data
"""

import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Paths
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, 'public', 'data')
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')

# Create output directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

def load_json_data(filename):
    """Load JSON data from public/data directory"""
    filepath = os.path.join(DATA_DIR, filename)
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)

def setup_document_styles(doc):
    """Configure document styles"""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

def add_header(doc, name, description_lines, contact_info):
    """Add resume header with name, description, and contact information"""
    # Name
    name_para = doc.add_paragraph(name)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.runs[0]
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    
    # Description (2-3 lines)
    for desc_line in description_lines:
        desc_para = doc.add_paragraph(desc_line)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.runs[0]
        desc_run.font.size = Pt(11)
        desc_run.font.name = 'Calibri'
    
    # Contact Info
    contact_para = doc.add_paragraph(contact_info)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.runs[0]
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Calibri'
    
    doc.add_paragraph()  # Spacing

def add_section_heading(doc, title):
    """Add a section heading"""
    para = doc.add_paragraph(title.upper())
    para_format = para.paragraph_format
    para_format.space_before = Pt(12)
    para_format.space_after = Pt(6)
    
    run = para.runs[0]
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.underline = True  # Use underline instead of border

def add_experience_section(doc, experiences):
    """Add professional experience section"""
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
    
    for exp in experiences:
        # Company and Role with duration on right
        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_after = Pt(3)
        
        # Role (Bold)
        role_run = header_para.add_run(exp['role'])
        role_run.font.bold = True
        role_run.font.size = Pt(11)
        role_run.font.name = 'Calibri'
        
        # Company
        header_para.add_run(f" | {exp['company']}")
        
        # Add tab for right alignment
        tab_stops = header_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        
        # Duration (Right aligned)
        duration_run = header_para.add_run(f"\t{exp['duration']}")
        duration_run.font.italic = True
        duration_run.font.size = Pt(10)
        duration_run.font.name = 'Calibri'
        
        # Responsibilities
        for responsibility in exp['responsibilities']:
            para = doc.add_paragraph(responsibility, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            
            run = para.runs[0]
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        
        doc.add_paragraph()  # Spacing between experiences

def add_projects_section(doc, projects):
    """Add projects section with only AI/ML projects"""
    add_section_heading(doc, "PROJECTS")
    
    for project in projects:
        # Project Name (Bold)
        name_para = doc.add_paragraph()
        name_para.paragraph_format.space_after = Pt(3)
        
        name_run = name_para.add_run(project['name'])
        name_run.font.bold = True
        name_run.font.size = Pt(11)
        name_run.font.name = 'Calibri'
        
        # Skills in parentheses (italic)
        if project.get('skills'):
            skills_text = ", ".join(project['skills'])
            skills_run = name_para.add_run(f" ({skills_text})")
            skills_run.font.italic = True
            skills_run.font.size = Pt(10)
            skills_run.font.name = 'Calibri'
        
        # Description Points
        for point in project.get('descriptionPoints', []):
            para = doc.add_paragraph(point, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            
            run = para.runs[0]
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        
        doc.add_paragraph()  # Spacing between projects

def add_skills_section(doc, skills_data):
    """Add technical skills section"""
    add_section_heading(doc, "TECHNICAL SKILLS")
    
    skill_categories = []
    for category, items in skills_data.items():
        if category in ["AI & Machine Learning", "Data Science & Analytics", 
                       "Image Processing & Computer Vision", "Languages"]:
            skill_categories.append((category, items))
    
    for category, items in skill_categories:
        skill_para = doc.add_paragraph()
        skill_para.paragraph_format.space_after = Pt(3)
        
        # Category heading (Bold)
        category_run = skill_para.add_run(f"{category}: ")
        category_run.font.bold = True
        category_run.font.size = Pt(10)
        category_run.font.name = 'Calibri'
        
        # Skills list
        skill_list = ", ".join(items)
        skills_run = skill_para.add_run(skill_list)
        skills_run.font.size = Pt(10)
        skills_run.font.name = 'Calibri'

def add_education_section(doc):
    """Add education section"""
    add_section_heading(doc, "EDUCATION")
    
    # M.Tech
    edu1_para = doc.add_paragraph()
    edu1_para.paragraph_format.space_after = Pt(3)
    
    edu1_run = edu1_para.add_run("Lovely Professional University – Ludhiana, Punjab")
    edu1_run.font.size = Pt(11)
    edu1_run.font.name = 'Calibri'
    
    # Set tab stop for right alignment
    tab_stops1 = edu1_para.paragraph_format.tab_stops
    tab_stops1.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Date right aligned
    date1_run = edu1_para.add_run(f"\tSept 2021 – July 2023")
    date1_run.font.italic = True
    date1_run.font.size = Pt(10)
    date1_run.font.name = 'Calibri'
    
    # Degree
    degree1_para = doc.add_paragraph()
    degree1_para.paragraph_format.space_after = Pt(6)
    degree1_run = degree1_para.add_run("M.Tech Computer Science; 9.08")
    degree1_run.font.size = Pt(10)
    degree1_run.font.name = 'Calibri'
    
    # B.Tech
    edu2_para = doc.add_paragraph()
    edu2_para.paragraph_format.space_after = Pt(3)
    
    edu2_run = edu2_para.add_run("NIT, Agartala – Agartala, Tripura")
    edu2_run.font.size = Pt(11)
    edu2_run.font.name = 'Calibri'
    
    # Set tab stop for right alignment
    tab_stops2 = edu2_para.paragraph_format.tab_stops
    tab_stops2.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Date right aligned
    date2_run = edu2_para.add_run(f"\tAug 2016 – June 2020")
    date2_run.font.italic = True
    date2_run.font.size = Pt(10)
    date2_run.font.name = 'Calibri'
    
    # Degree
    degree2_para = doc.add_paragraph()
    degree2_para.paragraph_format.space_after = Pt(6)
    degree2_run = degree2_para.add_run("B.Tech ECE; 8.01")
    degree2_run.font.size = Pt(10)
    degree2_run.font.name = 'Calibri'

def add_certifications_section(doc, certificates):
    """Add certifications section"""
    add_section_heading(doc, "CERTIFICATIONS")
    
    # Filter relevant certifications
    relevant_certs = [cert for cert in certificates if any(
        skill in cert.get('title', '').lower() or 
        any(s in ['Deep Learning', 'Data Analytics', 'Machine Learning', 'Data Science'] 
            for s in cert.get('skills', []))
        for skill in ['deep learning', 'data', 'machine learning', 'ml', 'ai']
    )]
    
    for cert in relevant_certs[:5]:  # Top 5 most relevant
        cert_para = doc.add_paragraph()
        cert_para.paragraph_format.space_after = Pt(2)
        
        # Title (Bold)
        title_run = cert_para.add_run(cert['title'])
        title_run.font.bold = True
        title_run.font.size = Pt(10)
        title_run.font.name = 'Calibri'
        
        # Issuer and Date
        issuer_text = f" | {cert['issuer']}"
        if cert.get('issuedOn'):
            issuer_text += f" | {cert['issuedOn']}"
        
        issuer_run = cert_para.add_run(issuer_text)
        issuer_run.font.size = Pt(10)
        issuer_run.font.name = 'Calibri'

def generate_resume():
    """Main function to generate the resume"""
    print("Loading data...")
    
    # Load data
    projects_data = load_json_data('projects.json')
    experience_data = load_json_data('experience.json')
    skills_data = load_json_data('skills.json')
    certificates_data = load_json_data('certificates.json')
    
    # Filter AI/ML projects only
    ai_projects = [p for p in projects_data if p.get('category') == 'AI and Machine Learning']
    
    # Select top 3 projects (prioritize diverse types)
    selected_projects = []
    project_types = {}
    
    for project in ai_projects:
        project_type = 'NLP' if 'NER' in project['name'] or 'NLP' in project.get('skills', []) else \
                      'Computer Vision' if 'Vision' in project['name'] or 'Computer Vision' in project.get('skills', []) else \
                      'Time Series' if 'Forecasting' in project['name'] else 'Other'
        
        if project_type not in project_types:
            project_types[project_type] = []
        project_types[project_type].append(project)
    
    # Select diverse projects
    selected_projects = []
    if 'NLP' in project_types:
        selected_projects.append(project_types['NLP'][0])  # Medical NER
    if 'Computer Vision' in project_types:
        selected_projects.append(project_types['Computer Vision'][0])  # Retail or School
    if 'Time Series' in project_types:
        selected_projects.append(project_types['Time Series'][0])  # Forecasting
    
    # Fill remaining slots
    remaining = 3 - len(selected_projects)
    for project in ai_projects:
        if project not in selected_projects and remaining > 0:
            selected_projects.append(project)
            remaining -= 1
    
    print(f"Selected {len(selected_projects)} AI/ML projects:")
    for p in selected_projects:
        print(f"  - {p['name']}")
    
    # Create document
    doc = Document()
    setup_document_styles(doc)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Contact Information
    # TODO: Replace [Your Portfolio URL] with your actual portfolio website URL
    contact_info = "Email: tarunsharma080997@gmail.com | Phone: +91 9774587305 | LinkedIn: linkedin.com/in/tarun-kumar-988887159 | Portfolio: [Your Portfolio URL]"
    
    # Description lines
    description_lines = [
        "Data Sciences Analyst with expertise in AI, Machine Learning, and Data Science.",
        "Specializing in computer vision, image processing, and building intelligent systems that solve real-world problems.",
        "Experienced in developing production-ready ML models for object detection, forecasting, and information extraction."
    ]
    
    # Add sections
    add_header(doc, "TARUN KUMAR", description_lines, contact_info)
    add_skills_section(doc, skills_data)
    add_experience_section(doc, experience_data)
    add_projects_section(doc, selected_projects)
    add_certifications_section(doc, certificates_data)
    add_education_section(doc)
    
    # Save document
    timestamp = datetime.now().strftime("%Y%m%d")
    output_filename = f"Tarun_Kumar_Resume_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    doc.save(output_path)
    print(f"\nResume generated successfully!")
    print(f"Output: {output_path}")
    
    return output_path

if __name__ == "__main__":
    try:
        generate_resume()
    except Exception as e:
        print(f"Error generating resume: {e}")
        import traceback
        traceback.print_exc()

